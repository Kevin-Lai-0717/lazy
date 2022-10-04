from selenium import webdriver
from selenium.webdriver.common.by import By
import re
import time
import datetime
import pyperclip as pc
from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt, Cm
import os
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class Insight:
    def __init__(self, account, password):
        self.account = account
        self.password = password

    def open_chrome(self):
        self.driver = webdriver.Chrome("chromedriver.exe")

    def connect(self):
        options = webdriver.ChromeOptions()
        options.add_argument('ignore-certificate-errors')
        self.driver = webdriver.Chrome(options=options)
        self.driver.get("https://172.23.11.100/login")
        time.sleep(2)
        self.driver.find_element(By.ID, "username").send_keys(self.account)
        self.driver.find_element(By.ID, "password").send_keys(self.password)
        time.sleep(1)
        self.driver.find_element(By.ID, "authform").click()
        time.sleep(5)

    def get_data(self):
        today = time.strftime("%#m/%#d", time.localtime())
        list1 = []
        # -----爬取 剩餘空間-----
        text = self.driver.find_element(By.XPATH, "//*[@id='ext-comp-1015']").text
        for word in text.splitlines(keepends=True):
            new_word = str(word)
            new_word = re.sub("\n", "", new_word).replace("TiB", "TB")
            list1.append(new_word)
        # -----爬取 Datastore Usage-----
        text2 = self.driver.find_element(By.ID, "app_status_ds").text
        self.new_text2 = text2.split(" ")

        # -----複製結果-----
        self.text = today + " ：" + list1[3] + " & "
        pc.copy(self.text)
        self.check = re.sub("%", "", self.new_text2[2])

        print("=======================================================")
        print(self.text + self.new_text2[2])
        print("結果已複製到剪貼簿")
        self.driver.close()

    def edit_text_to_docx(self):
        weekday = datetime.date.today().weekday()

        self.docx = Document(r'C:\Users\zxc578623\Desktop\工作紀錄.docx')
        # -----新增段落與文字-----
        p = self.docx.add_paragraph("")
        p.add_run(self.text).font.size = Pt(14)
        # -----檢查數據使用率有沒有>95-----
        if int(self.check) >= 95:
            run = p.add_run(self.new_text2[2])
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.font.size = Pt(14)

        else:
            p.add_run(self.new_text2[2]).font.size = Pt(14)
        # -----如果今天是星期五-----
        if weekday == 4:
            self.docx.add_page_break()

        self.docx.save(r'C:\Users\zxc578623\Desktop\工作紀錄.docx')

    def add_pictures(self):
        self.docx = Document(r'C:\Users\zxc578623\Desktop\工作紀錄.docx')
        # -----測試的話記得加
        # self.docx.add_page_break()
        today = datetime.date.today()
        startdate = today + datetime.timedelta(days=-4)
        # -----處理圖檔資料夾
        imgdirpath = r"C:\Users\zxc578623\Documents\\"
        imgdir = os.listdir(imgdirpath)
        imgdir.sort(key=lambda fn: os.path.getmtime(imgdirpath + fn))
        print(imgdir)
        s = -1
        dir_new = os.path.join(imgdirpath, imgdir[s])
        print(dir_new)

        # -----判斷是不是指定的資料夾
        try:
            while True:
                if os.path.isdir(dir_new) == False:
                    s -= 1
                    dir_new = os.path.join(imgdirpath, imgdir[s])
                    continue
                else:
                    if imgdir[s].startswith("file-download-") == False:
                        s -= 1
                        continue
                    else:
                        if imgdir[s].endswith(".zip"):
                            s -= 1
                            continue
                        else:
                            break
        except:
            print("沒有指定的圖檔資料夾")
            os._exit(0)

        # -----處理圖檔格式
        imgpath = dir_new + "\\"
        imglist = os.listdir(imgpath)
        imglist.sort(key=lambda x: int(x.replace("IMG_", "").split(".")[0]))
        s = 0
        for week in range(5):
            # -----第一行+入日期
            p = self.docx.add_paragraph("")
            p.add_run(startdate.strftime("%#m/%#d")).font.size = Pt(14)
            startdate += datetime.timedelta(days=1)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p = self.docx.add_paragraph("")
            p.add_run("左側冷氣：").font.size = Pt(14)

            # -----插入圖片
            pic = self.docx.add_picture(imgpath + imglist[s])
            s += 1
            pic.width = Cm(18.44)
            pic.height = Cm(24.58)

            # -----第二頁
            for i in range(3):
                self.docx.add_paragraph("")
            pic = self.docx.add_picture(imgpath + imglist[s])
            s += 1
            pic.width = Cm(18.44)
            pic.height = Cm(24.58)

            # -----第三頁
            self.docx.add_paragraph("")
            p = self.docx.add_paragraph("")
            p.add_run("右側冷氣：").font.size = Pt(14)
            pic = self.docx.add_picture(imgpath + imglist[s])
            s += 1
            pic.width = Cm(18.44)
            pic.height = Cm(24.58)

            # -----第四頁
            for i in range(3):
                self.docx.add_paragraph("")
            pic = self.docx.add_picture(imgpath + imglist[s])
            s += 1
            pic.width = Cm(18.44)
            pic.height = Cm(24.58)

            # -----第五頁
            self.docx.add_paragraph("")
            p = self.docx.add_paragraph("")
            p.add_run("機櫃正面：").font.size = Pt(14)
            pic = self.docx.add_picture(imgpath + imglist[s])
            s += 1
            pic.width = Cm(18.44)
            pic.height = Cm(24.58)
        self.docx.save(r'C:\Users\zxc578623\Desktop\工作紀錄-完成.docx')

